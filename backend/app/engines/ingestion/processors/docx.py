import io
from docx import Document
from ..base import BaseSourceProcessor, ExtractionResult
from ..validators import validate_file


class DOCXProcessor(BaseSourceProcessor):
    source_type = "docx"

    def validate(self, filename: str = "", content: bytes = b"", **kwargs) -> None:
        validate_file(filename, content)

    def extract(self, filename: str = "", content: bytes = b"", filepath: str = "", **kwargs) -> ExtractionResult:
        if filepath:
            doc = Document(filepath)
        else:
            doc = Document(io.BytesIO(content))

        sections = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            if "Heading" in style:
                sections.append(f"## {text}")
            else:
                sections.append(text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                sections.append("\n".join(rows))

        full_text = "\n\n".join(sections)
        metadata = {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
        title = filename or "Word Document"
        return ExtractionResult(text=full_text, metadata=metadata, title=title)
